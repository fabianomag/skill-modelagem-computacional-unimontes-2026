#!/usr/bin/env python3
"""Create the clean academic DOCX starter used by the skill.

The generated file intentionally contains structure and named styles, not a
solved example.  Future runs replace the bracketed drafting prompts with the
new problem's verified content.
"""

from __future__ import annotations

import argparse
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


TEAM = tuple(f"[Integrante {index}]" for index in range(1, 6))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_fixed(table, widths_twips: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table_pr = table._tbl.tblPr
    layout = table_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_twips:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_twips, strict=True):
            set_cell_width(cell, width)


def set_paragraph_box(paragraph, fill: str, border: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border)
    borders.append(left)
    p_pr.append(borders)
    spacing = p_pr.get_or_add_spacing()
    spacing.set(qn("w:before"), "120")
    spacing.set(qn("w:after"), "120")


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0x24, 0x37, 0x46)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(26)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x17, 0x32, 0x4D)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

    heading_specs = {
        "Heading 1": (16, RGBColor(0x2E, 0x74, 0xB5)),
        "Heading 2": (13, RGBColor(0x17, 0x32, 0x4D)),
        "Heading 3": (11.5, RGBColor(0x0F, 0x76, 0x6E)),
    }
    for name, (size, color) in heading_specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(5)

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor(0x52, 0x5F, 0x6B)

    if "Equation Block" not in styles:
        equation = styles.add_style("Equation Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        equation = styles["Equation Block"]
    equation.font.name = "Cambria Math"
    equation.font.size = Pt(11)
    equation.paragraph_format.left_indent = Mm(5)
    equation.paragraph_format.right_indent = Mm(5)
    equation.paragraph_format.space_before = Pt(6)
    equation.paragraph_format.space_after = Pt(6)

    for name, color in (
        ("Decision Callout", RGBColor(0x0B, 0x62, 0x3E)),
        ("Limit Callout", RGBColor(0x9B, 0x1C, 0x1C)),
    ):
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.left_indent = Mm(5)
        style.paragraph_format.right_indent = Mm(5)


def add_cover(document: Document, title: str, course: str, institution: str) -> None:
    for _ in range(3):
        document.add_paragraph()
    eyebrow = document.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = eyebrow.add_run(course.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

    paragraph = document.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(title)

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Relatório de modelagem, comparação e decisão")

    document.add_paragraph()
    line = document.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.add_run("Equipe").bold = True
    for member in TEAM:
        member_paragraph = document.add_paragraph()
        member_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        member_paragraph.add_run(member)

    for _ in range(2):
        document.add_paragraph()
    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(institution).bold = True
    footer.add_run(f"\n{datetime.now().year}")
    footer.add_run().add_break(WD_BREAK.PAGE)


def add_data_table(document: Document) -> None:
    table = document.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    set_table_fixed(table, [2200, 2200, 2200])
    headers = ("Entrada", "Saída observada", "Unidades")
    values = (
        ("[nome/símbolo]", "[nome/símbolo]", "[unidades]"),
        ("[valor ou faixa]", "[valor ou faixa]", "[observação]"),
    )
    for index, label in enumerate(headers):
        cell = table.cell(0, index)
        cell.text = label
        set_cell_shading(cell, "17324D")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row_index, row_values in enumerate(values, start=1):
        for col_index, value in enumerate(row_values):
            table.cell(row_index, col_index).text = value
    document.add_paragraph("Tabela 1 — Dados fornecidos e respectivas unidades.", style="Caption")


def create_document(
    title: str,
    course: str,
    institution: str,
    include_computational_appendix: bool = False,
) -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.72)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)

    configure_styles(document)
    add_cover(document, title, course, institution)

    header = section.header.paragraphs[0]
    header.text = course
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Modelagem Computacional · ")
    add_page_field(footer)
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    document.add_heading("Resumo e decisão", level=1)
    document.add_paragraph(
        "[Resuma o problema, o método, a principal evidência verificada e o limite da "
        "recomendação. Escreva este resumo após concluir as análises.]"
    )
    callout = document.add_paragraph("[Decisão final somente após os testes.]", style="Decision Callout")
    set_paragraph_box(callout, "E8F5EE", "0B623E")

    document.add_heading("1. Problema e dados", level=1)
    document.add_paragraph("[Reproduza fielmente o enunciado e identifique entrada, saída e domínio.]")
    add_data_table(document)

    document.add_heading("2. Método", level=1)
    document.add_paragraph(
        "[Explique conceitualmente o ajuste, a função de erro, os resíduos e as métricas antes da forma matricial.]"
    )
    equation = document.add_paragraph("[equação central e definição de símbolos]", style="Equation Block")
    set_paragraph_box(equation, "EEF2F5", "2E74B5")

    document.add_heading("3. Modelos candidatos", level=1)
    document.add_heading("3.1 Um subtítulo por candidato", level=2)
    document.add_paragraph(
        "[Duplique esta subseção para cada família realmente comparada. Em cada uma, apresente forma, "
        "parâmetros, hipótese e resultado verificado. Não force o problema a ter apenas dois candidatos.]"
    )

    document.add_heading("4. Resultados e diagnóstico", level=1)
    document.add_paragraph(
        "[Mostre previsões, resíduos, métricas, gráficos e ao menos uma conferência manual por candidato.]"
    )

    document.add_heading("5. Comparação e escolha do modelo", level=1)
    document.add_paragraph(
        "[Combine aderência ao problema, resíduos, métricas, complexidade e domínio. Não escolha por uma métrica isolada.]"
    )

    document.add_heading("6. Aplicação solicitada", level=1)
    document.add_paragraph("[Use o modelo escolhido apenas dentro das condições sustentadas pelos dados.]")

    document.add_heading("7. Limitações e conclusão", level=1)
    document.add_paragraph("[Responda objetivamente às perguntas do enunciado.]")
    limit = document.add_paragraph(
        "[Declare faixa válida, número de observações, incertezas e limites de extrapolação.]",
        style="Limit Callout",
    )
    set_paragraph_box(limit, "FDECEC", "9B1C1C")

    if include_computational_appendix:
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_heading("Apêndice computacional", level=1)
        document.add_paragraph(
            "[Inclua rastreabilidade dos cálculos ou referência ao código exigido pela rubrica ou pela reprodutibilidade.]"
        )

    properties = document.core_properties
    now = datetime.now(timezone.utc)
    properties.title = title
    properties.subject = "Relatório acadêmico de Modelagem Computacional"
    properties.author = ""
    properties.keywords = "modelagem computacional, ajuste, resíduos, comparação de modelos"
    properties.created = now
    properties.modified = now
    properties.comments = "Template gerado pela skill modelagem-computacional-grupo."
    return document


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("output", type=Path)
    parser.add_argument("--title", default="[Título do problema]")
    parser.add_argument("--course", default="Modelagem Computacional")
    parser.add_argument("--institution", default="[Instituição]")
    parser.add_argument("--include-computational-appendix", action="store_true")
    args = parser.parse_args()

    output = args.output.expanduser().resolve()
    if output.suffix.lower() != ".docx":
        parser.error("A saída deve usar a extensão .docx")
    output.parent.mkdir(parents=True, exist_ok=True)
    document = create_document(
        args.title,
        args.course,
        args.institution,
        include_computational_appendix=args.include_computational_appendix,
    )
    document.save(output)
    print(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
